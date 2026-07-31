"""
ATS Scout — Recruitment Intelligence Platform
Flask Backend Server
Run: python app.py
"""

import os
import json
import uuid
import tempfile
import urllib.request
import urllib.error
import ssl
from flask import Flask, request, jsonify, render_template, send_from_directory
from werkzeug.utils import secure_filename
import pdfplumber
from docx import Document as DocxDocument

# ─────────────────────────────────────────────────────────────────────────────
#  APP SETUP
# ─────────────────────────────────────────────────────────────────────────────
app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
#  HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(filepath):
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return '\n'.join(text_parts).strip()
    except Exception as e:
        raise ValueError(f"Could not read PDF: {str(e)}")


def extract_text_from_docx(filepath):
    """Extract text from DOCX using python-docx."""
    try:
        doc = DocxDocument(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return '\n'.join(paragraphs).strip()
    except Exception as e:
        raise ValueError(f"Could not read DOCX: {str(e)}")


def extract_text_from_txt(filepath):
    """Read plain text file."""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file — try saving as UTF-8.")


def extract_text(filepath, extension):
    """Route file to correct parser."""
    ext = extension.lower()
    if ext == 'pdf':
        return extract_text_from_pdf(filepath)
    elif ext in ('docx', 'doc'):
        return extract_text_from_docx(filepath)
    elif ext == 'txt':
        return extract_text_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


def call_anthropic(api_key, prompt, max_tokens=2500):
    """Call Anthropic Claude API using urllib (no SDK required)."""
    payload = json.dumps({
        "model": "claude-sonnet-4-6",
        "max_tokens": max_tokens,
        "messages": [{"role": "user", "content": prompt}]
    }).encode('utf-8')

    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        method='POST'
    )
    req.add_header('Content-Type', 'application/json')
    req.add_header('x-api-key', api_key)
    req.add_header('anthropic-version', '2023-06-01')

    # Create SSL context that doesn't verify (for corporate networks)
    ctx = ssl.create_default_context()
    try:
        with urllib.request.urlopen(req, context=ctx, timeout=60) as resp:
            data = json.loads(resp.read().decode('utf-8'))
            return data['content'][0]['text']
    except urllib.error.HTTPError as e:
        error_body = e.read().decode('utf-8')
        try:
            error_data = json.loads(error_body)
            raise ValueError(error_data.get('error', {}).get('message', f'API error {e.code}'))
        except json.JSONDecodeError:
            raise ValueError(f'API error {e.code}: {error_body[:200]}')
    except urllib.error.URLError as e:
        raise ValueError(f'Network error: {str(e.reason)}. Check internet connection.')


def build_prompt(jd_text, candidate_name, resume_text):
    """Build the ATS evaluation prompt."""
    # Truncate to avoid token limits
    jd = jd_text[:3000] if len(jd_text) > 3000 else jd_text
    resume = resume_text[:6000] if len(resume_text) > 6000 else resume_text

    return f"""You are a strict, expert ATS screening system specialized in Embedded Systems engineering recruitment.

Evaluate this candidate resume against the job description. Be precise, objective, and thorough.

=== JOB DESCRIPTION ===
{jd}

=== CANDIDATE: {candidate_name} ===
{resume}

Respond ONLY with a valid JSON object (no markdown, no explanation, just raw JSON):

{{
  "name": "{candidate_name}",
  "atsScore": <integer 0-100>,
  "roleMatch": "<Testing|Development|Integration|Mixed|None>",
  "roleMatchDetail": "<1 sentence on role fit>",
  "mandatory": "<Pass|Fail>",
  "mandatoryDetails": {{
    "education": "<Pass|Fail>",
    "embeddedC": "<Pass|Fail>",
    "basicC": "<Pass|Fail>",
    "debugging": "<Pass|Fail>"
  }},
  "scoringBreakdown": {{
    "educationMatch": <0-10>,
    "embeddedCBasicC": <0-25>,
    "roleSpecificSkills": <0-25>,
    "relevantProjects": <0-10>,
    "relevantExperience": <0-15>,
    "debuggingExperience": <0-10>,
    "resumeQuality": <0-5>
  }},
  "matchedSkills": ["skill1", "skill2"],
  "missingSkills": ["skill1", "skill2"],
  "matchedKeywords": [
    {{"keyword": "Embedded C", "importance": "Critical", "found": true, "impact": "+15"}},
    {{"keyword": "Linux", "importance": "Critical", "found": false, "impact": "-10"}},
    {{"keyword": "Python", "importance": "High", "found": false, "impact": "-7"}},
    {{"keyword": "Git", "importance": "High", "found": true, "impact": "+5"}},
    {{"keyword": "RTOS", "importance": "High", "found": false, "impact": "-5"}},
    {{"keyword": "Debugging/RCA", "importance": "High", "found": true, "impact": "+5"}},
    {{"keyword": "Shell Scripting", "importance": "Medium", "found": false, "impact": "-3"}},
    {{"keyword": "Jenkins/CI-CD", "importance": "Medium", "found": false, "impact": "-3"}},
    {{"keyword": "Device Drivers", "importance": "High", "found": false, "impact": "-5"}},
    {{"keyword": "B.Tech/BE", "importance": "Critical", "found": true, "impact": "+10"}}
  ],
  "projects": [
    {{"name": "project name", "relevance": "<Excellent|Good|Average|Weak>", "description": "why relevant or not"}}
  ],
  "strengths": ["strength 1", "strength 2", "strength 3"],
  "weaknesses": ["weakness 1", "weakness 2"],
  "rejectionRisks": ["specific risk 1", "specific risk 2"],
  "acceptanceReasons": ["reason 1 why this candidate should be accepted", "reason 2"],
  "rejectionReasons": ["reason 1 why this candidate should be rejected", "reason 2"],
  "feedback": {{
    "positive": ["observation 1", "observation 2"],
    "negative": ["concern 1", "concern 2"],
    "suggestions": ["suggestion 1", "suggestion 2"]
  }},
  "interviewReadiness": "<High|Medium|Low>",
  "recommendation": "<Strongly Recommended|Recommended|Consider|Reject>",
  "hrNotes": "2-sentence HR summary for this candidate.",
  "yearsExperience": "<0-1|1-3|3-5|5+|Unknown>",
  "education": "highest qualification found"
}}

SCORING WEIGHTS (must sum to atsScore):
- Education (B.Tech/BE): 10pts
- Embedded C + Basic C: 25pts
- Role Skills (Testing=Python+Auto, Dev=Linux+Drivers, Integr=Git+Jenkins): 25pts
- Relevant Projects: 10pts
- Experience: 15pts
- Debugging/RCA: 10pts
- Resume Quality: 5pts

Scoring thresholds:
>=85: Strongly Recommended | 70-84: Recommended | 55-69: Consider | <55: Reject

Return ONLY valid JSON."""


# ─────────────────────────────────────────────────────────────────────────────
#  ROUTES
# ─────────────────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/parse-file', methods=['POST'])
def parse_file():
    """Parse uploaded file and return extracted text."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': f'Unsupported file type. Use PDF, DOCX, DOC, or TXT'}), 400

    try:
        filename = secure_filename(file.filename)
        ext = filename.rsplit('.', 1)[1].lower()
        # Save temporarily
        tmp_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{uuid.uuid4().hex}.{ext}')
        file.save(tmp_path)

        try:
            text = extract_text(tmp_path, ext)
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

        if not text or len(text) < 50:
            return jsonify({'error': 'File appears empty or unreadable. Try copying text manually.'}), 400

        # Derive candidate name from filename
        name = filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').title()
        return jsonify({
            'success': True,
            'name': name,
            'text': text,
            'charCount': len(text),
            'fileType': ext.upper()
        })

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500


@app.route('/api/evaluate', methods=['POST'])
def evaluate():
    """Evaluate a single candidate resume against JD."""
    data = request.get_json()

    api_key = data.get('apiKey', '').strip()
    jd_text = data.get('jdText', '').strip()
    candidate_name = data.get('candidateName', 'Candidate').strip()
    resume_text = data.get('resumeText', '').strip()

    if not api_key:
        return jsonify({'error': 'API key is required'}), 400
    if not jd_text:
        return jsonify({'error': 'Job description is required'}), 400
    if not resume_text:
        return jsonify({'error': 'Resume text is required'}), 400
    if len(resume_text) < 50:
        return jsonify({'error': 'Resume text is too short'}), 400

    try:
        prompt = build_prompt(jd_text, candidate_name, resume_text)
        raw = call_anthropic(api_key, prompt)

        # Strip markdown wrappers if present
        clean = raw.strip()
        if clean.startswith('```'):
            clean = clean.split('```', 2)[-1] if clean.count('```') >= 2 else clean
            clean = clean.replace('json', '', 1).strip()
            if clean.endswith('```'):
                clean = clean[:-3].strip()

        result = json.loads(clean)
        return jsonify({'success': True, 'result': result})

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except json.JSONDecodeError as e:
        return jsonify({'error': f'AI returned invalid response format. Please try again. ({str(e)[:100]})'}), 500
    except Exception as e:
        return jsonify({'error': f'Evaluation failed: {str(e)}'}), 500


@app.route('/api/validate-key', methods=['POST'])
def validate_key():
    """Quick check if API key format is valid."""
    data = request.get_json()
    key = data.get('apiKey', '').strip()
    if not key:
        return jsonify({'valid': False, 'message': 'No key provided'})
    if not key.startswith('sk-ant'):
        return jsonify({'valid': False, 'message': 'Invalid format — key must start with sk-ant'})
    return jsonify({'valid': True, 'message': 'Key format looks good'})


@app.route('/health')
def health():
    return jsonify({'status': 'ok', 'version': '2.0.0'})


if __name__ == '__main__':
    print("=" * 60)
    print("  ATS Scout — Recruitment Intelligence Platform")
    print("  Version 2.0.0")
    print("=" * 60)
    print(f"  Server : http://localhost:5000")
    print(f"  Open your browser and go to: http://localhost:5000")
    print("=" * 60)
    app.run(debug=True, host='0.0.0.0', port=5000)
